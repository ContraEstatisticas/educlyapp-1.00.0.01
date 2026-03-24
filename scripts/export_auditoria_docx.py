from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from markdown import markdown

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_SOFT = RGBColor(0x48, 0x6A, 0x85)
TEXT = RGBColor(0x25, 0x2B, 0x33)
MUTED = RGBColor(0x6B, 0x72, 0x7A)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_border_bottom(paragraph, color: str = "D9E1F2", size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field_run(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.8)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = ACCENT

    subtitle = styles.add_style("Audit Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = styles["Normal"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    eyebrow = styles.add_style("Audit Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
    eyebrow.base_style = styles["Normal"]
    eyebrow.font.name = "Calibri"
    eyebrow.font.size = Pt(9.5)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = ACCENT_SOFT
    eyebrow.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_after = Pt(6)

    for style_name, size, color in (
        ("Heading 1", 18, ACCENT),
        ("Heading 2", 13.5, ACCENT_SOFT),
        ("Heading 3", 11.5, ACCENT),
    ):
        heading = styles[style_name]
        heading.font.name = "Georgia"
        heading.font.bold = True
        heading.font.size = Pt(size)
        heading.font.color.rgb = color
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True

    code_block = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_block.base_style = styles["Normal"]
    code_block.font.name = "Consolas"
    code_block.font.size = Pt(9.5)
    code_block.paragraph_format.left_indent = Cm(0.6)
    code_block.paragraph_format.right_indent = Cm(0.4)
    code_block.paragraph_format.space_before = Pt(4)
    code_block.paragraph_format.space_after = Pt(6)
    code_block.paragraph_format.line_spacing = 1.0

    quote = styles.add_style("Audit Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote.base_style = styles["Normal"]
    quote.font.italic = True
    quote.font.color.rgb = MUTED
    quote.paragraph_format.left_indent = Cm(0.6)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(6)


def configure_header_footer(document: Document, metadata: dict[str, str]) -> None:
    section = document.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_parts = ["Educly", metadata["header_label"], metadata["title"]]
    if metadata["date"]:
        header_parts.append(metadata["date"])
    run = p.add_run(" | ".join(part for part in header_parts if part))
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = p.add_run("Pagina ")
    prefix.font.name = "Calibri"
    prefix.font.size = Pt(8.5)
    prefix.font.color.rgb = MUTED
    add_field_run(p, "PAGE")


def humanize_stem(stem: str) -> str:
    return stem.replace("-", " ").replace("_", " ").strip().title()


def extract_metadata(input_path: Path, markdown_text: str) -> dict[str, str]:
    project = ""
    reference_date = ""
    title = ""

    for line in markdown_text.splitlines():
        if line.startswith("Data base da auditoria:"):
            reference_date = line.split(":", 1)[1].strip()
        if line.startswith("Data de criacao:"):
            reference_date = line.split(":", 1)[1].strip()
        if line.startswith("Projeto auditado:"):
            project = line.split(":", 1)[1].strip().strip("`")
        if line.startswith("# ") and not title:
            title = line[2:].strip()

    if not title:
        title = humanize_stem(input_path.stem)

    relative_source = str(input_path)
    try:
        relative_source = str(input_path.relative_to(Path.cwd()))
    except ValueError:
        relative_source = str(input_path)

    is_audit = "auditoria" in input_path.stem.lower() or "auditoria" in title.lower()
    is_runbook = "runbooks" in input_path.parts

    if is_audit:
        eyebrow = "AUDITORIA TECNICA"
        subtitle = "Seguranca, performance, backup e controle de acesso"
        intro = (
            "Linha de base da primeira auditoria completa do banco do Educly, "
            "consolidando seguranca, performance, backup e controle de acesso."
        )
        header_label = "Auditoria"
        date_label = "Data-base da auditoria"
        subject = "Auditoria de banco do Educly"
    elif is_runbook:
        eyebrow = "RUNBOOK OPERACIONAL"
        subtitle = "Procedimento, governanca e resposta operacional"
        intro = (
            "Documento operacional versionado do Educly, preparado para execucao, "
            "governanca e evidencias futuras da auditoria."
        )
        header_label = "Runbook"
        date_label = "Data de referencia"
        subject = f"Runbook operacional do Educly: {title}"
    else:
        eyebrow = "DOCUMENTO TECNICO"
        subtitle = "Documento versionado do Educly"
        intro = "Documento gerado automaticamente a partir do markdown versionado do projeto."
        header_label = "Documento"
        date_label = "Data de referencia"
        subject = f"Documento tecnico do Educly: {title}"

    return {
        "title": title,
        "subtitle": subtitle,
        "eyebrow": eyebrow,
        "intro": intro,
        "header_label": header_label,
        "date_label": date_label,
        "project": project or "Educly",
        "date": reference_date or "",
        "source": relative_source,
        "subject": subject,
    }


def add_cover(document: Document, metadata: dict[str, str]) -> None:
    document.add_paragraph(metadata["eyebrow"], style="Audit Eyebrow")

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(metadata["title"])

    subtitle = document.add_paragraph(style="Audit Subtitle")
    subtitle.add_run(metadata["subtitle"])

    intro = document.add_paragraph(style="Audit Subtitle")
    intro.add_run(metadata["intro"])
    intro.paragraph_format.space_after = Pt(20)

    meta_table = document.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_table.autofit = True

    rows = [
        ("Projeto", metadata["project"]),
        (metadata["date_label"], metadata["date"] or "Nao informado"),
        ("Documento-fonte", metadata["source"]),
    ]

    for idx, (label, value) in enumerate(rows):
        left = meta_table.rows[idx].cells[0]
        right = meta_table.rows[idx].cells[1]
        left.text = label
        right.text = value
        set_cell_shading(left, "DCE6F1")
        set_cell_shading(right, "F8FAFC")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = ACCENT

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note_run = note.add_run(
        "Versao DOCX gerada a partir do baseline versionado em Markdown."
    )
    note_run.italic = True
    note_run.font.color.rgb = MUTED

    document.add_page_break()

    toc_title = document.add_paragraph("Sumario", style="Heading 1")
    toc_title.paragraph_format.space_before = Pt(0)
    toc = document.add_paragraph()
    add_field_run(toc, 'TOC \\o "1-3" \\h \\z \\u')
    hint = document.add_paragraph(style="Audit Quote")
    hint.add_run(
        "Se o sumario nao aparecer preenchido de imediato, abra o arquivo no Word "
        "e atualize os campos do documento."
    )
    document.add_page_break()


def append_text_run(paragraph, text: str, *, bold=False, italic=False, code=False, link_href: str | None = None) -> None:
    if not text:
        return

    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic

    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(9.2)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

    if link_href:
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True


def append_inline(paragraph, node, *, bold=False, italic=False) -> None:
    if isinstance(node, NavigableString):
        append_text_run(paragraph, str(node), bold=bold, italic=italic)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()

    if name == "br":
        paragraph.add_run().add_break()
        return

    if name in {"strong", "b"}:
        for child in node.children:
            append_inline(paragraph, child, bold=True or bold, italic=italic)
        return

    if name in {"em", "i"}:
        for child in node.children:
            append_inline(paragraph, child, bold=bold, italic=True or italic)
        return

    if name == "code":
        append_text_run(paragraph, node.get_text(), bold=bold, italic=italic, code=True)
        return

    if name == "a":
        append_text_run(
            paragraph,
            node.get_text(),
            bold=bold,
            italic=italic,
            link_href=node.get("href"),
        )
        href = node.get("href")
        if href and href != node.get_text():
            append_text_run(paragraph, f" ({href})", italic=True)
        return

    for child in node.children:
        append_inline(paragraph, child, bold=bold, italic=italic)


def li_primary_children(li: Tag) -> Iterable[Tag | NavigableString]:
    for child in li.children:
        if isinstance(child, Tag) and child.name in {"ul", "ol"}:
            continue
        yield child


def add_list(document: Document, list_tag: Tag, level: int = 0) -> None:
    ordered = list_tag.name == "ol"
    style_name = "List Number" if ordered else "List Bullet"

    for li in list_tag.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.left_indent = Cm(0.55 * level)
        paragraph.paragraph_format.space_after = Pt(2)

        for child in li_primary_children(li):
            append_inline(paragraph, child)

        for nested in li.find_all(["ul", "ol"], recursive=False):
            add_list(document, nested, level + 1)


def add_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        rows = table_tag.find_all("tr")
    if not rows:
        return

    max_cols = max(len(row.find_all(["th", "td"], recursive=False) or row.find_all(["th", "td"])) for row in rows)
    table = document.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"

    for row_idx, row_tag in enumerate(rows):
        row = table.add_row().cells
        cells = row_tag.find_all(["th", "td"], recursive=False) or row_tag.find_all(["th", "td"])
        for col_idx in range(max_cols):
            cell = row[col_idx]
            if col_idx < len(cells):
                cell.text = cells[col_idx].get_text(" ", strip=True)
            else:
                cell.text = ""

            if row_idx == 0:
                set_cell_shading(cell, "DCE6F1")
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = ACCENT
            else:
                set_cell_shading(cell, "F8FAFC")

    document.add_paragraph()


def add_code_block(document: Document, code_text: str) -> None:
    for line in code_text.rstrip().splitlines() or [""]:
        paragraph = document.add_paragraph(style="Code Block")
        set_paragraph_border_bottom(paragraph, color="E2E8F0", size="4")
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)

    document.add_paragraph()


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def render_markdown(document: Document, markdown_text: str) -> None:
    html = markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html, "html.parser")

    seen_main_title = False
    seen_item_section = False
    skip_heading_level: int | None = None

    for node in soup.contents:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue

        tag_name = node.name.lower()

        if tag_name == "h1":
            if not seen_main_title:
                seen_main_title = True
                continue

        if tag_name in {"h1", "h2", "h3", "h4"}:
            level = {"h1": 1, "h2": 1, "h3": 2, "h4": 3}[tag_name]
            text = clean_text(node.get_text(" ", strip=True))

            if skip_heading_level is not None:
                if level <= skip_heading_level:
                    skip_heading_level = None
                else:
                    continue

            if text == "O que ficou pendente":
                skip_heading_level = level
                continue

            if tag_name == "h2" and text.startswith("Item "):
                if seen_item_section:
                    document.add_page_break()
                seen_item_section = True

            paragraph = document.add_paragraph(style=f"Heading {level}")
            append_inline(paragraph, node)
            continue

        if skip_heading_level is not None:
            continue

        if tag_name == "p":
            paragraph = document.add_paragraph(style="Normal")
            append_inline(paragraph, node)
            continue

        if tag_name in {"ul", "ol"}:
            add_list(document, node)
            continue

        if tag_name == "blockquote":
            paragraph = document.add_paragraph(style="Audit Quote")
            for child in node.children:
                append_inline(paragraph, child)
            continue

        if tag_name == "pre":
            code = node.get_text()
            add_code_block(document, code)
            continue

        if tag_name == "table":
            add_table(document, node)
            continue

        if tag_name == "hr":
            paragraph = document.add_paragraph()
            set_paragraph_border_bottom(paragraph)
            continue


def export_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    markdown_text = input_path.read_text(encoding="utf-8")
    metadata = extract_metadata(input_path, markdown_text)

    document = Document()
    configure_document(document)
    configure_header_footer(document, metadata)

    core = document.core_properties
    core.author = "OpenAI Codex"
    core.title = metadata["title"]
    core.subject = metadata["subject"]
    core.comments = f"Arquivo DOCX gerado automaticamente a partir de {metadata['source']}."

    add_cover(document, metadata)
    render_markdown(document, markdown_text)

    if output_path.parent:
        output_path.parent.mkdir(parents=True, exist_ok=True)

    document.save(output_path)


def main() -> int:
    if len(sys.argv) not in {2, 3}:
        print("Uso: python scripts/export_auditoria_docx.py <input.md> [output.docx]")
        return 1

    input_path = Path(sys.argv[1]).resolve()
    output_path = (
        Path(sys.argv[2]).resolve()
        if len(sys.argv) == 3
        else input_path.with_suffix(".docx")
    )

    export_markdown_to_docx(input_path, output_path)
    print(f"DOCX gerado em: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
